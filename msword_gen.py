import docx
#from pylatexenc.latex2mathml import convert_latex_to_mathml
# Create a new document


def add_math_to_word(file_name, eqns):

    # for now just generating new file

    doc = docx.Document()

    for eqn in eqns:

        # Convert the LaTeX to MathML
        #math_ml = convert_latex_to_mathml(latex)

        # Insert the MathML text into Word
        run = doc.add_paragraph().add_run(text=eqn)
        run.font.math = True
        run.style.quick_style = True

        # Save the document


    doc.save(file_name)
